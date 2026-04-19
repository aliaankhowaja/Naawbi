"""
gen_d3.py — Generate D3.docx for Naawbi LMS Deliverable 3
CS3009 Software Engineering | Spring 2026
Run: python gen_d3.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ── Colours ─────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)   # #1F3864  headings / titles
GRAY   = RGBColor(0x60, 0x60, 0x60)   # #606060  subtitles / captions
LBLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # #D6E4F0  table header fill
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT = r"E:\codebases\uni\Naawbi\docs\D3\D3.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set cell background shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_hr(doc, color="1F3864", thickness=6):
    """Add a full-width horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def add_short_hr(doc, width_pct=30, color="1F3864"):
    """Centered short decorative rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def para(doc, text="", size=11, bold=False, italic=False,
         color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, font_name="Times New Roman"):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name  = font_name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p


def heading1(doc, text, numbered=True):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    return p


def heading2(doc, text):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    return p


def heading3(doc, text):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_doc_margins(section, top=1.0, bottom=1.0, left=1.2, right=1.0):
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


def add_header_footer(section, doc):
    """Add running page header with line below."""
    header = section.header
    header.is_linked_to_previous = False

    htable = header.add_table(1, 2, width=Inches(6.5))
    htable.style = "Table Grid"
    # remove all borders from the table itself
    for row in htable.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)

    lc = htable.cell(0, 0)
    rc = htable.cell(0, 1)
    lp = lc.paragraphs[0]
    rp = rc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lr = lp.add_run("Naawbi LMS — Deliverable 3")
    lr.font.name  = "Calibri"
    lr.font.size  = Pt(9)
    lr.font.color.rgb = NAVY

    rr = rp.add_run("CS3009 Spring 2026")
    rr.font.name  = "Calibri"
    rr.font.size  = Pt(9)
    rr.font.color.rgb = GRAY

    # Add a bottom border to the header paragraph (acts as the line)
    pPr  = lp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3864")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("ibbi1020  •  wajih-rathore  •  aliaankhowaja")
    fr.font.name  = "Calibri"
    fr.font.size  = Pt(9)
    fr.font.color.rgb = GRAY


def make_table(doc, headers, rows, small_font=False, landscape_note=False):
    """Create a bordered table with light-blue header row."""
    font_size = 9 if small_font else 10
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "D6E4F0")
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name  = "Calibri"
        run.font.bold  = True
        run.font.size  = Pt(font_size)
        run.font.color.rgb = NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_run(p, str(cell_text), size=font_size)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_inline_run(paragraph, text: str, size=10):
    """Add a run to a paragraph, handling **bold** and `code` inline markup."""
    # Split by bold (**text**) and code (`text`)
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(min(size, 10))
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(size)
    return paragraph


def add_body_para(doc, text: str, italic=False, gray=False, code_block=False, indent=False):
    """Add a body paragraph with inline markup support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)

    if code_block:
        run = p.add_run(text)
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        p.paragraph_format.left_indent = Inches(0.4)
        return p

    if italic and gray:
        run = p.add_run(text)
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(10)
        run.font.italic = True
        run.font.color.rgb = GRAY
        return p

    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold  = True
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(11)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.font.italic = True
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(11)
        else:
            run = p.add_run(part)
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(11)
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + 0.25 * level)

    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold  = True
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(11)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
        else:
            run = p.add_run(part)
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(11)
    return p


def add_code_block(doc, lines: list):
    """Add a shaded code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.4)
        set_para_shade(p, "F2F2F2")
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def set_para_shade(p, hex_color):
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


# ── Cover Page ───────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    set_doc_margins(section, top=1.2, bottom=1.0, left=1.2, right=1.0)

    # University name — small caps effect via all-caps
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("National University of Computer and Emerging Sciences")
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.all_caps = True
    run.font.color.rgb = NAVY

    # Course line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    run2 = p2.add_run("CS3009 — Software Engineering  |  Spring 2026")
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(10)
    run2.font.color.rgb = GRAY

    # Horizontal rule
    add_hr(doc, color="1F3864", thickness=8)

    # Spacer
    para(doc, space_before=14, space_after=0)

    # "Naawbi" — big navy title
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(4)
    run3 = p3.add_run("Naawbi")
    run3.font.name  = "Calibri"
    run3.font.size  = Pt(36)
    run3.font.bold  = True
    run3.font.color.rgb = NAVY

    # LMS subtitle
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after  = Pt(10)
    run4 = p4.add_run("Learning Management System")
    run4.font.name  = "Calibri"
    run4.font.size  = Pt(14)
    run4.font.color.rgb = GRAY

    # Short HR
    add_short_hr(doc)

    # Spacer
    para(doc, space_before=6, space_after=0)

    # "Deliverable 3"
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(4)
    p5.paragraph_format.space_after  = Pt(4)
    run5 = p5.add_run("Deliverable 3")
    run5.font.name  = "Calibri"
    run5.font.size  = Pt(22)
    run5.font.bold  = True
    run5.font.color.rgb = NAVY

    # Deliverable subtitle
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Pt(0)
    p6.paragraph_format.space_after  = Pt(20)
    run6 = p6.add_run("Requirements & Sprint 3 Documentation")
    run6.font.name  = "Calibri"
    run6.font.size  = Pt(13)
    run6.font.color.rgb = GRAY

    # Metadata table — borderless, right-label / left-value
    meta = [
        ("Team Name",       "Ibwaan"),
        ("Team Lead",       "Ibraheem Farooq"),
        ("Members",         "Ibraheem Farooq, Wajih-Ur-Raza Asif, Ali Aan"),
        ("Submission Date", "20 April 2026"),
        ("Repository",      "github.com/aliaankhowaja/Naawbi"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders from metadata table
    tbl_el  = tbl._tbl
    tbl_pr  = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    for i, (label, value) in enumerate(meta):
        lc = tbl.cell(i, 0)
        vc = tbl.cell(i, 1)

        # Remove cell borders too
        for cell in [lc, vc]:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label + ":")
        lr.font.name  = "Calibri"
        lr.font.size  = Pt(11)
        lr.font.bold  = True
        lr.font.color.rgb = NAVY

        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run("  " + value)
        vr.font.name  = "Calibri"
        vr.font.size  = Pt(11)
        vr.font.color.rgb = BLACK

    # Cover page footer
    para(doc, space_before=30, space_after=0)
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(6)
    rf = pf.add_run("ibbi1020  •  wajih-rathore  •  aliaankhowaja")
    rf.font.name  = "Calibri"
    rf.font.size  = Pt(10)
    rf.font.color.rgb = GRAY

    doc.add_page_break()


# ── Table of Contents ────────────────────────────────────────────────────────

TOC_ENTRIES = [
    ("PART A — Project Team Info", 0),
    ("Team Updates", 1),
    ("PART B — Development Methodology: Scrum", 0),
    ("1. Sprint 3 Overview", 1),
    ("2. Sprint 3 Backlog", 1),
    ("2a. Handoff from Sprint 2", 2),
    ("2b. User Stories Selected for Sprint 3", 2),
    ("2c. Sub User Stories and Task Breakdowns", 2),
    ("3. Structured Specifications (Sprint 3)", 1),
    ("Spec 1: US 1–4 — Authentication", 2),
    ("Spec 2: US 5 — Role-Based Access Control", 2),
    ("Spec 3: US 12–13 — Announcement Attachments", 2),
    ("Spec 4: US 15–16 — Assignments Tab and People Tab", 2),
    ("Spec 5: US 19–20 — Instructor Grading Interface", 2),
    ("Spec 6: US 21–22 — Gradebook Views", 2),
    ("4. Scrum Board (Trello Snapshots)", 1),
    ("5. Sprint 3 Burn-Down Chart", 1),
    ("6. Testing Activities for Sprint 3", 1),
    ("6a. Software Test Plan", 2),
    ("6b. Black-Box Test Cases", 2),
    ("6c. White-Box Testing", 2),
    ("6d. Bug Report", 2),
    ("6e. Test Execution Results", 2),
    ("PART C — Deliverable 3 Document", 0),
    ("1. Introduction", 1),
    ("2. Project Vision", 1),
    ("3. Intended Use of the System", 1),
    ("4. Features and Overall Functionality", 1),
    ("5. User Stories (Sprint 3)", 1),
    ("6. Non-Functional Requirements (NFR) — Final State", 1),
    ("7. Testing and Quality Assurance", 1),
    ("8. Sprint 3 Implementation Screenshots", 1),
    ("8. Work Division for Sprint 3", 1),
    ("9. Retrospective", 1),
    ("11. Final System Summary", 1),
]


def build_toc(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(10)
    rt = p_title.add_run("Table of Contents")
    rt.font.name  = "Calibri"
    rt.font.size  = Pt(16)
    rt.font.bold  = True
    rt.font.color.rgb = NAVY

    add_hr(doc, color="1F3864", thickness=4)
    para(doc, space_before=4, space_after=0)

    for entry, level in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        indent = Inches(0.3 * level)
        p.paragraph_format.left_indent = indent
        run = p.add_run(entry)
        if level == 0:
            run.font.bold  = True
            run.font.size  = Pt(11)
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(6)
        elif level == 1:
            run.font.size  = Pt(10.5)
            run.font.color.rgb = NAVY
        else:
            run.font.size  = Pt(10)
            run.font.color.rgb = GRAY
        run.font.name = "Calibri"

    doc.add_page_break()


# ── PART A ───────────────────────────────────────────────────────────────────

def build_part_a(doc):
    p_part = doc.add_paragraph()
    p_part.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_part.paragraph_format.space_before = Pt(0)
    p_part.paragraph_format.space_after  = Pt(6)
    rp = p_part.add_run("PART A — Project Team Info")
    rp.font.name  = "Calibri"
    rp.font.size  = Pt(18)
    rp.font.bold  = True
    rp.font.color.rgb = NAVY
    add_hr(doc)

    heading2(doc, "Team Updates")
    add_body_para(doc, "No changes to team membership or roles. Ibraheem Farooq continues as Team Lead.")

    make_table(doc,
        headers=["Member", "Role(s)", "GitHub"],
        rows=[
            ["Ibraheem Farooq", "PM / Scrum Master / Developer", "github.com/ibbi1020"],
            ["Wajih-Ur-Raza Asif", "Requirements Architect / Tester / Developer", "github.com/wajih-rathore"],
            ["Ali Aan Khowaja", "Developer / UI Designer", "github.com/aliaankhowaja"],
        ]
    )

    add_body_para(doc,
        "The team agreement established in Deliverable 1 remains fully active. "
        "No changes to communication methods, meeting schedules, version control practices, "
        "work division, submission responsibilities, or contingency planning were needed during Sprint 3."
    )

    doc.add_page_break()


# ── PART B ───────────────────────────────────────────────────────────────────

def build_part_b(doc):
    p_part = doc.add_paragraph()
    p_part.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_part.paragraph_format.space_before = Pt(0)
    p_part.paragraph_format.space_after  = Pt(6)
    rp = p_part.add_run("PART B — Development Methodology: Scrum")
    rp.font.name  = "Calibri"
    rp.font.size  = Pt(18)
    rp.font.bold  = True
    rp.font.color.rgb = NAVY
    add_hr(doc)

    # ── Section 1: Sprint 3 Overview ─────────────────────────────────────────
    heading1(doc, "1. Sprint 3 Overview")
    make_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Sprint Duration",         "26 March 2026 – 20 April 2026"],
            ["Sprint Goal",
             "Close the platform's remaining feature gaps: implement user authentication with "
             "registration and login, enforce role-based access control, add file and link "
             "attachments to announcements, build the Assignments and People tabs in the course "
             "home, and deliver the full instructor grading interface and student-facing gradebook."],
            ["Carried from D2 Backlog", "US 1–4 (Authentication), US 12–13 (Announcement Attachments), US 19–22 (Grading & Gradebook)"],
            ["Added in Sprint 3",       "US 15–16 (Course Home: Assignments Tab and People Tab)"],
        ]
    )

    # ── Section 2: Sprint 3 Backlog ───────────────────────────────────────────
    heading1(doc, "2. Sprint 3 Backlog")

    heading2(doc, "2a. Handoff from Sprint 2")
    add_body_para(doc,
        "Deliverable 2's final Trello snapshot showed the following user stories in the product "
        "backlog at sprint close:"
    )
    add_bullet(doc, "**US 1–4** — Authentication (register, login, role assignment, logout)")
    add_bullet(doc, "**US 12–13** — Announcement attachments (file and link)")
    add_bullet(doc, "**US 19–22** — Grading interface and gradebook (instructor and student views)")
    add_body_para(doc, "During Sprint 3 planning, the team also added:")
    add_bullet(doc,
        "**US 15–16** — Course home navigation (Assignments tab with status badges; People tab with "
        "course roster) — these were partially implied by the D2 assignment workflow but were never "
        "formally scoped or built."
    )

    heading2(doc, "2b. User Stories Selected for Sprint 3")
    us_stories = [
        ("US 1", "As a User, I want to register with a username, email, password, and role so that I have my own personalised account on the platform."),
        ("US 2", "As a User, I want to log in with my email and password so that I can access my course workspace."),
        ("US 3", "As a User, I want to be assigned a role (Instructor or Student) at registration so that the system tailors my experience and enforces appropriate permissions."),
        ("US 4", "As a User, I want to log out so that my session is terminated and my account is secure."),
        ("US 5", "As the System, I want to enforce role-based access control so that instructors and students only see and use the features appropriate to their role."),
        ("US 12", "As an Instructor, I want to attach a file to an announcement so that I can share reference materials directly from the course stream."),
        ("US 13", "As an Instructor, I want to attach a URL to an announcement so that I can share links to external resources without requiring a file upload."),
        ("US 15 (Assignments Tab)", "As a User, I want a dedicated Assignments tab in the course home showing all assignments and their status at a glance, so I do not have to navigate away from the course to see my workload."),
        ("US 16 (People Tab)", "As a User, I want a People tab in the course home showing everyone enrolled in the course with their role, so I know who is in my class."),
        ("US 19", "As an Instructor, I want to open a grading view for an assignment so that I can see all student submissions in one place."),
        ("US 20", "As an Instructor, I want to assign a numeric grade and written feedback to each submission so that students receive clear, actionable results."),
        ("US 21", "As a Student, I want to see my grades and instructor feedback for each graded assignment so that I know how I performed."),
        ("US 22", "As an Instructor, I want a class-wide gradebook grid showing all students and assignments so that I have a complete overview of course performance without opening submissions individually."),
    ]
    for us_id, us_text in us_stories:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(2)
        rb = p.add_run(us_id + " — ")
        rb.font.bold  = True
        rb.font.name  = "Times New Roman"
        rb.font.size  = Pt(11)
        rt2 = p.add_run(us_text)
        rt2.font.name  = "Times New Roman"
        rt2.font.size  = Pt(11)

    heading2(doc, "2c. Sub User Stories and Task Breakdowns")

    sub_sections = [
        ("US 1–4 — Authentication (Login + Registration)", [
            "Design tabbed LoginView with Login and Register tabs switching between two forms",
            "Register tab: username field, email field, password field, role dropdown (Instructor / Student), Register button",
            "Login tab: email and password fields, Sign In button, inline error label for failed attempts",
            "`PasswordUtil.hash()` — SHA-256 hash applied to password before storage and comparison (never stored plain)",
            "`User.register()` — inserts user record; returns empty Optional if username or email already taken",
            "`LoginController.handleLogin()` — compares hashed input against stored hash; populates `Session` singleton on success",
            "`LoginController.handleRegister()` — validates all fields non-empty; calls `User.register()`; navigates to catalog on success",
            "`Session` singleton — holds id, username, email, role; cleared on logout",
            "`CourseCatalogController.handleLogout()` — clears Session, reloads LoginView into the primary Stage",
            "Log Out button wired into CourseCatalogView header",
        ]),
        ("US 5 — Role-Based Access Control", [
            "`CourseCatalogController.initialize()` — queries `course_enrollments` for students; queries `courses.created_by` for instructors; filters catalog list accordingly",
            "Instructor-only action guard — checks `Session.getInstance().getRole()` before opening Create Assignment, Post Announcement, or Grade Submissions; shows Access Denied alert for students",
            "Course Stream tab: Create Announcement button visible only to instructors",
            "Assignments tab: Grade Submissions button shown only to instructors; status badge shown only to students",
        ]),
        ("US 12–13 — Announcement Attachments", [
            "`CreateAnnouncementController` extended with two optional attachment sections:",
            "  File attachment: FileChooser opens system file picker; selected file path and filename stored with `is_link = false`",
            "  Link attachment: TextField for URL input; stored with `is_link = true`, `link_url` populated",
            "`announcement_attachments` table: `id`, `announcement_id`, `file_name`, `link_url`, `is_link`",
            "Attachment chips rendered below announcement body in the course stream; file chip opens the file, link chip opens the URL in browser",
        ]),
        ("US 15 — Assignments Tab", [
            "`Assignment.fetchWithStatusForUser(courseId, userId)` — LEFT JOIN on submissions; derives status per assignment: Submitted, Late, Missing, or Not Submitted based on deadline and submission timestamp",
            "Assignments tab in `CourseCatalogController` renders one card per assignment with: title, deadline, point value, late-policy badge",
            "Students see a colour-coded status badge; instructors see a Grade Submissions button that opens the Grading modal",
        ]),
        ("US 16 — People Tab", [
            "`Course.fetchEnrolledUsers(courseId)` — queries `course_enrollments` JOIN `users`; returns list of [username, email, role]",
            "People tab renders role-badged roster cards separated into Instructor and Student sections; each card shows avatar initial, username, email, and role label",
        ]),
        ("US 19–20 — Instructor Grading Interface", [
            "`grade` (NUMERIC) and `feedback` (TEXT) columns added to `submissions` table in `DB.createTables()`",
            "`Submission.fetchByAssignmentId(assignmentId)` — loads all submissions for an assignment",
            "`Submission.saveGrade(submissionId, grade, feedback)` — parameterised UPDATE on the submissions row",
            "`GradingView.fxml` — two-panel modal: left ListView of submissions (student name + status), right panel with grade input field, feedback TextArea, and Save button",
            "`GradingController` — loads submissions on open; pre-fills grade and feedback if already graded; validates grade is numeric and <= total_points before saving",
        ]),
        ("US 21–22 — Gradebook Views", [
            "`Submission.fetchGradedForStudent(courseId, userId)` — returns all graded submissions for one student in a course",
            "`Submission.fetchGradebookForCourse(courseId)` — returns all student-assignment-grade triples for the full class",
            "`GradebookView.fxml` + `GradebookController` — role-adaptive:",
            "  Student mode: table with assignment name, score / total_points, and feedback per row",
            "  Instructor mode: GridPane with enrolled students as rows, assignments as columns, numeric grade or — per cell",
            "Grades tab added to course home tab bar; opens GradebookView as a modal",
        ]),
    ]

    for title, bullets in sub_sections:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(8)
        p_sub.paragraph_format.space_after  = Pt(2)
        rsub = p_sub.add_run(title)
        rsub.font.name  = "Calibri"
        rsub.font.size  = Pt(11)
        rsub.font.bold  = True
        rsub.font.color.rgb = NAVY
        for b in bullets:
            add_bullet(doc, b)

    # ── Section 3: Structured Specifications ─────────────────────────────────
    heading1(doc, "3. Structured Specifications (Sprint 3)")

    specs = [
        ("Spec 1: US 1–4 — Authentication (Register + Login + Logout)", [
            ("Actor",            "Any user (Student or Instructor)"),
            ("Inputs",           "Register: username, email, password, role (Instructor / Student). Login: email, password."),
            ("Processing",
             "Registration: validate all fields non-empty; hash password with SHA-256; call `User.register()` which inserts user and returns empty Optional on duplicate username/email. "
             "Login: hash submitted password; compare against stored hash; if match, load user into `Session` singleton. "
             "Logout: call `Session.logout()` and navigate back to LoginView."),
            ("Outputs",          "On successful register or login: user lands on CourseCatalogView with session active. On logout: session cleared; user lands on LoginView."),
            ("Preconditions",    "Register: username and email not already in `users` table. Login: account exists with a hashed password."),
            ("Postconditions",   "Session holds user id, username, email, role. Logout leaves Session empty; no credentials persist."),
            ("Exceptions",
             'Wrong password or unknown email: "Incorrect email or password" inline label. '
             'Empty fields: "Please fill in all fields". '
             'Duplicate username/email on register: "Username or email is already taken".'),
        ]),
        ("Spec 2: US 5 — Role-Based Access Control", [
            ("Actor",            "System (enforced at every protected point)"),
            ("Inputs",           "`Session.getInstance().getRole()` — set at login/register"),
            ("Processing",
             "On catalog load: students see only courses where they have a `course_enrollments` record; instructors see courses where `courses.created_by = userId`. "
             "On instructor-only action (Post Announcement, Create Assignment, Grade Submissions): check role and show Access Denied alert if role is `student`. "
             "Instructor-only UI elements (Create Announcement button, Grade Submissions button) are conditionally rendered."),
            ("Outputs",          "Students see only their enrolled courses. Instructor-only features are inaccessible to students both visually and via guard logic."),
            ("Preconditions",    "User is logged in; Session is populated."),
            ("Postconditions",   "No student can trigger an instructor action even by direct controller call."),
            ("Exceptions",       "No unauthenticated code path exists in the application after Sprint 3; LoginView is always the first screen."),
        ]),
        ("Spec 3: US 12–13 — Announcement Attachments", [
            ("Actor",            "Instructor (attach); all enrolled users (view)"),
            ("Inputs",           "Optional: file selected via system FileChooser OR URL typed into a text field. Both are optional — an announcement with no attachment is valid."),
            ("Processing",
             "File: store original filename and file path in `announcement_attachments` with `is_link = false`. "
             "URL: store in `link_url` with `is_link = true`. "
             "Both records are linked to the announcement by foreign key. A single announcement can carry one file and one link attachment independently."),
            ("Outputs",          "Attachment chips rendered inline below announcement body in the course stream. File chip opens the file; link chip opens the URL."),
            ("Preconditions",    "Instructor is creating or has created an announcement."),
            ("Postconditions",   "Attachment record persisted; visible to all enrolled students on next stream load."),
            ("Exceptions",       "No file/URL provided: announcement saves normally with no attachment record. Cleared attachment fields are ignored."),
        ]),
        ("Spec 4: US 15–16 — Assignments Tab and People Tab", [
            ("Actor",            "Student and Instructor"),
            ("Inputs",           "Current course selection in catalog"),
            ("Processing",
             "Assignments tab: `Assignment.fetchWithStatusForUser(courseId, userId)` — LEFT JOIN `submissions`; status derived as Submitted (before deadline), Late (after deadline with late policy), Missing (past deadline, no submission), Not Submitted (deadline not reached, no submission). "
             "People tab: `Course.fetchEnrolledUsers(courseId)` — JOIN `users`; separate into Instructor and Student sections."),
            ("Outputs",
             'Assignments: cards with title, deadline, point value, late-policy badge, and status badge (students) or "Grade Submissions" button (instructors). '
             "People: roster cards with avatar initial, username, email, and role badge."),
            ("Preconditions",    "A course is selected; user is enrolled or is the course instructor."),
            ("Postconditions",   "Both tabs refresh on every course selection event."),
            ("Exceptions",       'No assignments: empty state label "No assignments yet." No enrolled users: empty state label.'),
        ]),
        ("Spec 5: US 19–20 — Instructor Grading Interface", [
            ("Actor",            "Instructor"),
            ("Inputs",           "Selected submission from ListView; numeric grade (required); feedback text (optional)"),
            ("Processing",
             "On open: load all submissions for the assignment via `Submission.fetchByAssignmentId()`; pre-fill grade/feedback if a previous grade exists. "
             "On Save: validate grade is numeric and within 0–total_points; call `Submission.saveGrade(submissionId, grade, feedback)` which issues a parameterised UPDATE."),
            ("Outputs",          "`submissions.grade` and `submissions.feedback` updated. ListView item reflects graded state. Changes immediately visible in student gradebook."),
            ("Preconditions",    "Instructor is logged in and owns the course; assignment exists."),
            ("Postconditions",   "Submission record has non-null `grade` and optionally non-null `feedback`."),
            ("Exceptions",       "No submissions: empty ListView with empty state label. Non-numeric grade input: inline validation error; Save is blocked."),
        ]),
        ("Spec 6: US 21–22 — Gradebook Views", [
            ("Actor",            "Student (personal grades) and Instructor (class gradebook)"),
            ("Inputs",           "Current course ID; logged-in user role from Session"),
            ("Processing",
             "Student: `Submission.fetchGradedForStudent(courseId, userId)` — returns all graded submissions for this student; renders one row per assignment with score, total_points, and feedback. "
             "Instructor: `Submission.fetchGradebookForCourse(courseId)` — returns all student-assignment-grade triples; builds a GridPane with students as rows and assignments as columns."),
            ("Outputs",          "Student: personal grade table with score and feedback per assignment. Instructor: class-wide grid with numeric grades or — for ungraded submissions."),
            ("Preconditions",    "Course is selected; user is enrolled (student) or owns the course (instructor)."),
            ("Postconditions",   "Read-only display — no mutations happen from either Gradebook view."),
            ("Exceptions",       'No graded submissions: "No grades yet." empty state for student. Empty grid with "—" cells for instructor if nothing has been graded.'),
        ]),
    ]

    for spec_title, spec_rows in specs:
        heading2(doc, spec_title)
        make_table(doc,
            headers=["Field", "Detail"],
            rows=[[label, detail] for label, detail in spec_rows]
        )

    # ── Section 4: Scrum Board ────────────────────────────────────────────────
    heading1(doc, "4. Scrum Board (Trello Snapshots)")
    add_body_para(doc,
        "[Snapshot 1 — Sprint 3 backlog at sprint start: US 1–5, 12–13, 15–16, 19–22 in Backlog column]",
        italic=True, gray=True
    )
    add_body_para(doc,
        "[Snapshot 2 — Mid-sprint: US 1–5 and US 12–13 in Done; US 15–16 and US 19–22 in In Progress]",
        italic=True, gray=True
    )
    add_body_para(doc,
        "[Snapshot 3 — Sprint end: all stories Done; product backlog empty]",
        italic=True, gray=True
    )

    # ── Section 5: Burn-Down Chart ────────────────────────────────────────────
    heading1(doc, "5. Sprint 3 Burn-Down Chart")
    add_body_para(doc,
        "[Burn-down chart: planned story points vs. completed story points across sprint duration, 26 Mar – 20 Apr 2026]",
        italic=True, gray=True
    )

    # ── Section 6: Testing ────────────────────────────────────────────────────
    heading1(doc, "6. Testing Activities for Sprint 3")

    # 6a Test Plan
    heading2(doc, "6a. Software Test Plan")
    make_table(doc,
        headers=["Field", "Detail"],
        rows=[
            ["Scope",
             "Sprint 3 features: Authentication (register/login/logout), RBAC, Announcement Attachments, Assignments Tab, People Tab, Grading Interface, Gradebook"],
            ["Objectives",
             "Verify each feature meets its acceptance criteria; detect regression in Sprint 1/2 features; validate RBAC enforcement across all protected actions"],
            ["Test Types",
             "Black-box functional testing (equivalence partitioning, boundary value analysis, error guessing); white-box structural testing (branch coverage on key controller methods)"],
            ["Test Environment",
             "Windows 11, Java 24, JavaFX 21, PostgreSQL 16 — seeded with data/seed.sql (3 users: ali = instructor, ibbi + sara = students; 2 courses; 6 assignments; 6 submissions in mixed states)"],
            ["Entry Criteria",
             "`make seed` succeeds; `make compile` produces zero errors; application launches to LoginView"],
            ["Exit Criteria",
             "All P1 (critical) test cases pass; no unfixed P1 bugs; all identified P2 bugs documented with status"],
            ["Tester", "Wajih-Ur-Raza Asif"],
        ]
    )

    # 6b Black-Box Test Cases — note: small font for these wide tables
    heading2(doc, "6b. Black-Box Test Cases")

    bb_modules = [
        ("Module: Authentication — Login", [
            ["TC-01", "Valid partition", "Email: ali@naawbi.edu, Password: password123", "Login succeeds; CourseCatalogView loads; session role = instructor", "Pass"],
            ["TC-02", "Valid partition", "Email: ibbi@naawbi.edu, Password: password123", "Login succeeds; CourseCatalogView loads; session role = student", "Pass"],
            ["TC-03", "Invalid partition", "Email: ali@naawbi.edu, Password: wrongpass", 'Inline error: "Incorrect email or password"; stays on LoginView', "Pass"],
            ["TC-04", "Invalid partition", "Email: nobody@naawbi.edu, Password: password123", 'Inline error: "Incorrect email or password"', "Pass"],
            ["TC-05", 'Boundary / empty', 'Email: "", Password: ""', 'Inline error: "Please fill in all fields"', "Pass"],
            ["TC-06", "Error guessing", "Email: ' OR 1=1 --, Password: anything", "Login fails — parameterised query prevents SQL injection", "Pass"],
        ]),
        ("Module: Authentication — Registration", [
            ["TC-07", "Valid partition", "New unique username, email, password, role = Student", "Account created; redirected to CourseCatalogView", "Pass"],
            ["TC-08", "Invalid partition", "Username already taken", 'Error: "Username or email is already taken"', "Pass"],
            ["TC-09", "Invalid partition", "Email already taken", 'Error: "Username or email is already taken"', "Pass"],
            ["TC-10", "Boundary / empty", "Any field left blank", 'Error: "Please fill in all fields"', "Pass"],
        ]),
        ("Module: RBAC — Instructor-Only Guard", [
            ["TC-11", "Invalid partition", 'Student ibbi clicks "Create Assignment"', "Access Denied alert shown; no form opens", "Pass"],
            ["TC-12", "Valid partition", 'Instructor ali clicks "Create Assignment"', "Create Assignment form opens normally", "Pass"],
            ["TC-13", "Valid partition", "Student logs in", "Catalog shows only enrolled courses (CS101, CS202)", "Pass"],
            ["TC-14", "Valid partition", "Instructor logs in", "Catalog shows only courses created by ali", "Pass"],
        ]),
        ("Module: Grading Interface", [
            ["TC-15", "Valid partition", 'Grade = 85, total_points = 100, feedback = "Good work"', "Grade saved; submission shows as graded in ListView", "Pass"],
            ["TC-16", "Boundary value", "Grade = 0", "Grade saved as 0; valid", "Pass"],
            ["TC-17", "Boundary value", "Grade = total_points (e.g., 100)", "Grade saved at maximum; valid", "Pass"],
            ["TC-18", "Boundary value", "Grade = total_points + 1 (e.g., 101)", "Inline validation error; Save blocked", "Pass"],
            ["TC-19", "Invalid partition", 'Grade = "abc" (non-numeric)', "Inline validation error; Save blocked", "Pass"],
            ["TC-20", "Valid partition", "No feedback provided; grade = 50", "Grade saved with null feedback; valid", "Pass"],
        ]),
        ("Module: Submission Status Derivation", [
            ["TC-21", "Valid partition", "Submission timestamp before deadline", "Status badge: Submitted", "Pass"],
            ["TC-22", "Valid partition", "Submission timestamp after deadline; late policy = allowed", "Status badge: Late", "Pass"],
            ["TC-23", "Valid partition", "Deadline passed; no submission; late policy = allowed", "Status badge: Missing", "Pass"],
            ["TC-24", "Valid partition", "Deadline in future; no submission", "Status badge: Not Submitted", "Pass"],
            ["TC-25", "Invalid partition", "Deadline passed; late policy = hard deadline; attempt to submit", "Submission blocked", "Pass"],
        ]),
    ]

    for mod_title, tc_rows in bb_modules:
        heading3(doc, mod_title)
        make_table(doc,
            headers=["TC #", "Technique", "Input", "Expected Output", "Status"],
            rows=tc_rows,
            small_font=True
        )

    # 6c White-Box Testing
    heading2(doc, "6c. White-Box Testing")

    heading3(doc, "LoginController — handleLogin(): Branch Coverage")
    add_code_block(doc, [
        "Method branches:",
        "  B1: email or password field empty → show error, return",
        "  B2: User.findByEmail() returns empty Optional → show error, return",
        "  B3: stored hash does not match submitted hash → show error, return",
        "  B4: all checks pass → populate Session, navigate to catalog",
    ])
    make_table(doc,
        headers=["Branch", "Test Case", "Covered"],
        rows=[
            ["B1 — empty fields",   "TC-05",         "Yes"],
            ["B2 — unknown email",  "TC-04",         "Yes"],
            ["B3 — wrong password", "TC-03",         "Yes"],
            ["B4 — successful login","TC-01, TC-02", "Yes"],
        ]
    )
    p_cov = doc.add_paragraph()
    r_cov = p_cov.add_run("Branch coverage: 4/4 (100%)")
    r_cov.font.bold  = True
    r_cov.font.name  = "Times New Roman"
    r_cov.font.size  = Pt(11)
    p_cov.paragraph_format.space_after = Pt(8)

    heading3(doc, "GradingController — handleSave(): Branch Coverage")
    add_code_block(doc, [
        "Method branches:",
        "  B1: grade field empty or non-numeric → show error, return",
        "  B2: grade > total_points → show error, return",
        "  B3: grade < 0 → show error, return",
        "  B4: all checks pass → call Submission.saveGrade(), update ListView",
    ])
    make_table(doc,
        headers=["Branch", "Test Case", "Covered"],
        rows=[
            ["B1 — non-numeric input",   "TC-19",              "Yes"],
            ["B2 — grade over maximum",  "TC-18",              "Yes"],
            ["B3 — grade below minimum", "Boundary: grade = -1 (manual)", "Yes"],
            ["B4 — valid save",          "TC-15, TC-16, TC-17","Yes"],
        ]
    )
    p_cov2 = doc.add_paragraph()
    r_cov2 = p_cov2.add_run("Branch coverage: 4/4 (100%)")
    r_cov2.font.bold  = True
    r_cov2.font.name  = "Times New Roman"
    r_cov2.font.size  = Pt(11)
    p_cov2.paragraph_format.space_after = Pt(8)

    heading3(doc, "Assignment.fetchWithStatusForUser() — Statement Coverage")
    add_body_para(doc,
        "The method executes a LEFT JOIN query and derives submission status via conditional logic. "
        "All four status paths (Submitted, Late, Missing, Not Submitted) are exercised by TC-21 through TC-24 "
        "using the seeded dataset, achieving statement coverage of all reachable branches in the status derivation block."
    )

    # 6d Bug Report
    heading2(doc, "6d. Bug Report")
    make_table(doc,
        headers=["Bug ID", "Module", "Description", "Severity", "Status"],
        rows=[
            ["BUG-01", "Grading",
             "GradingView showed a NullPointerException when opened for an assignment with zero submissions — fetchByAssignmentId() returned an empty list and the controller attempted to access index 0",
             "P1", "Fixed — added empty-list guard before ListView population"],
            ["BUG-02", "RBAC",
             "After merging feature/enrollment-rbac, instructor could see all courses including ones they did not create because the created_by field was not being set on course save",
             "P1", "Fixed — CreateCourseController updated to set created_by = Session.getInstance().getId()"],
            ["BUG-03", "Gradebook",
             "Instructor gradebook GridPane column headers were misaligned by one column when a student had no submissions for the first assignment",
             "P2", "Fixed — GridPane population rewritten to iterate all assignments independently of submission presence"],
            ["BUG-04", "Attachments",
             "Announcement attachment chip did not render after posting when the course stream was refreshed — AnnouncementAttachment.fetchByAnnouncementId() was not being called during stream load",
             "P1", "Fixed — stream loader updated to fetch and render attachments per announcement"],
        ],
        small_font=True
    )

    # 6e Test Execution Results
    heading2(doc, "6e. Test Execution Results")
    add_body_para(doc,
        "Testing was performed manually against the seeded database (`make seed`) using all three test accounts:"
    )
    make_table(doc,
        headers=["Account", "Role", "Scenarios Tested"],
        rows=[
            ["ali@naawbi.edu", "Instructor",
             "Login, RBAC guard, create assignment, post announcement with attachments, grade all submissions, view class gradebook, People tab, logout"],
            ["ibbi@naawbi.edu", "Student",
             "Login, register (new account), enrollment filtering, submit assignment, view submission status, personal gradebook, To-Do list, logout"],
            ["sara@naawbi.edu", "Student",
             "Login, second student perspective, verify gradebook shows sara's grades separately from ibbi's, People tab roster"],
        ]
    )
    add_body_para(doc,
        "All 25 black-box test cases passed. All 4 identified bugs (BUG-01 through BUG-04) were fixed before submission. "
        "No unfixed P1 bugs remain. One P2 issue was found and fixed (BUG-03). "
        "System behaviour matched acceptance criteria across all Sprint 3 user stories."
    )

    doc.add_page_break()


# ── PART C ───────────────────────────────────────────────────────────────────

def build_part_c(doc):
    p_part = doc.add_paragraph()
    p_part.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_part.paragraph_format.space_before = Pt(0)
    p_part.paragraph_format.space_after  = Pt(6)
    rp = p_part.add_run("PART C — Deliverable 3 Document")
    rp.font.name  = "Calibri"
    rp.font.size  = Pt(18)
    rp.font.bold  = True
    rp.font.color.rgb = NAVY
    add_hr(doc)

    # ── 1. Introduction ───────────────────────────────────────────────────────
    heading1(doc, "1. Introduction")
    heading2(doc, "Context")
    add_body_para(doc,
        "Naawbi is a university-level Learning Management System built iteratively by Team Ibwaan for "
        "CS3009 Software Engineering (Spring 2026). This document covers Deliverable 3, the output of "
        "Sprint 3 and the final sprint of the project."
    )

    heading2(doc, "Progress from Sprint 2 to Sprint 3")
    add_body_para(doc,
        "Sprint 2 delivered the content and submission layer — announcements, assignments, file-based submissions, "
        "submission status tracking, and the student To-Do list. All of that ran without any authentication; any user "
        "could reach any feature. Sprint 3 closes the platform: every user now registers with an identity and role, "
        "logs in to a session-guarded workspace, and participates in a complete grade cycle — from submission through "
        "instructor grading to a personal gradebook — while role-based access control prevents students from accessing "
        "instructor features and vice versa."
    )
    add_body_para(doc,
        "Sprint 3 also completed two features that were implied by D2's assignment workflow but were never formally "
        "built: the Assignments tab (giving students a per-course view of all assignments with submission status) and "
        "the People tab (showing the course roster with role badges)."
    )

    heading2(doc, "Carried from Sprint 2 Backlog")
    add_body_para(doc,
        "At the close of Sprint 2, the following user stories were explicitly listed in the product backlog "
        "(visible in the D2 Trello snapshot):"
    )
    make_table(doc,
        headers=["User Story", "Description"],
        rows=[
            ["US 1–4",  "Authentication — register, login, role assignment, logout"],
            ["US 12–13","Announcement attachments — file and link"],
            ["US 19–22","Grading interface and gradebook"],
        ]
    )
    add_body_para(doc,
        "Sprint 3 additionally scoped US 15–16 (Assignments tab + People tab) which had not appeared in any prior backlog."
    )

    # ── 2. Project Vision ─────────────────────────────────────────────────────
    heading1(doc, "2. Project Vision")
    add_body_para(doc,
        "Build a clean, reliable, course-centric LMS that instructors use to manage content and evaluations, "
        "and that students use to track progress and performance — without needing external tools for any core "
        "academic workflow."
    )
    add_body_para(doc,
        "Sprint 1 built the course structure. Sprint 2 populated it with content and submissions. Sprint 3 closes "
        "the loop: authenticated identities, enforced roles, grading, and a gradebook. The full feature set is now delivered."
    )

    # ── 3. Intended Use ───────────────────────────────────────────────────────
    heading1(doc, "3. Intended Use of the System")
    heading2(doc, "Stakeholders")
    make_table(doc,
        headers=["Stakeholder", "Primary Need", "Sprint 3 Impact"],
        rows=[
            ["Student",
             "Single place to track work, deadlines, and grades",
             "Logs in with a verified account; sees only enrolled courses; submits assignments; views per-course assignment status in the Assignments tab; sees personal grades and feedback in the Gradebook"],
            ["Instructor",
             "Unified workflow for content, collection, and evaluation",
             "Logs in; is the only role that can post, create assignments, and grade; can view a full class-wide gradebook at a glance; People tab shows course roster"],
            ["System",
             "Consistent enforcement of access rules",
             "RBAC enforced at every protected action via Session role check; no unauthenticated code path remains"],
        ]
    )

    # ── 4. Features ───────────────────────────────────────────────────────────
    heading1(doc, "4. Features and Overall Functionality")
    heading2(doc, "Cumulative Feature Summary (Deliverables 1–3)")
    make_table(doc,
        headers=["Feature", "User Stories", "Deliverable", "Status"],
        rows=[
            ["Course Creation",                           "US 5 (D1)",        "D1", "Done"],
            ["Course Catalogue and Home Page",            "US 6, US 9",       "D1", "Done"],
            ["Unique Class Code Enrollment",              "US 7, US 8",       "D1", "Done"],
            ["Course Announcements",                      "US 10, US 11",     "D2", "Done"],
            ["Assignment Creation with Policies",         "US 14, US 15 (D2)","D2", "Done"],
            ["File Submissions and Status Tracking",      "US 16, US 17, US 18","D2","Done"],
            ["Threaded Comments on Submissions",          "US 16 (extension)", "D2","Done"],
            ["To-Do List and Deadline Awareness",         "US 23, US 24",      "D2","Done"],
            ["User Registration and Login",               "US 1, US 2, US 3, US 4","D3","Done"],
            ["Logout",                                    "US 4",              "D3","Done"],
            ["Role-Based Access Control",                 "US 5 (D3)",         "D3","Done"],
            ["Announcement Attachments (File + Link)",    "US 12, US 13",      "D3","Done"],
            ["Assignments Tab with Status Badges",        "US 15 (D3)",        "D3","Done"],
            ["People Tab (Course Roster)",                "US 16 (D3)",        "D3","Done"],
            ["Instructor Grading Interface",              "US 19, US 20",      "D3","Done"],
            ["Student Personal Gradebook",                "US 21",             "D3","Done"],
            ["Instructor Class Gradebook",                "US 22",             "D3","Done"],
        ]
    )

    # ── 5. User Stories ───────────────────────────────────────────────────────
    heading1(doc, "5. User Stories (Sprint 3)")

    us_sections_c = [
        ("5.1 Authentication", [
            ("US 1", "As a User, I want to register with a username, email, password, and role so that I have my own account on the platform."),
            ("US 2", "As a User, I want to log in with my email and password so that I can access my personalised course workspace."),
            ("US 3", "As a User, I want to be assigned a role (Instructor or Student) at registration so that the system tailors my experience and enforces appropriate permissions."),
            ("US 4", "As a User, I want to log out so that my session is terminated and my account is secure."),
        ]),
        ("5.2 Role-Based Access Control", [
            ("US 5", "As the System, I want to enforce role-based access control so that instructors and students only see and use the features appropriate to their role; students see only enrolled courses and cannot trigger instructor-only actions."),
        ]),
        ("5.3 Announcement Attachments", [
            ("US 12", "As an Instructor, I want to attach a file to an announcement so that I can share reference materials directly from the course stream."),
            ("US 13", "As an Instructor, I want to attach a URL to an announcement so that I can share links to external resources without requiring a file upload."),
        ]),
        ("5.4 Course Home Navigation", [
            ("US 15", "As a User, I want a dedicated Assignments tab in the course home that shows all assignments with my submission status (students) or grading actions (instructors), so I can see my workload at a glance."),
            ("US 16", "As a User, I want a People tab in the course home that shows everyone enrolled in the course with their role badge, so I can see who is in my class."),
        ]),
        ("5.5 Grading", [
            ("US 19", "As an Instructor, I want to open a grading view for an assignment so that I can see all student submissions in one place."),
            ("US 20", "As an Instructor, I want to assign a numeric grade and written feedback to each submission so that students receive clear, actionable results."),
        ]),
        ("5.6 Gradebook", [
            ("US 21", "As a Student, I want to see my grades and instructor feedback for each graded assignment in a personal gradebook so that I know how I performed and where to improve."),
            ("US 22", "As an Instructor, I want a class-wide gradebook grid showing every student and every assignment so that I can track overall class performance without opening each submission individually."),
        ]),
    ]

    for sec_title, stories in us_sections_c:
        heading2(doc, sec_title)
        for us_id, us_text in stories:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(2)
            rb = p.add_run(us_id + " — ")
            rb.font.bold  = True
            rb.font.name  = "Times New Roman"
            rb.font.size  = Pt(11)
            rt2 = p.add_run(us_text)
            rt2.font.name  = "Times New Roman"
            rt2.font.size  = Pt(11)

    # ── 6. NFRs ───────────────────────────────────────────────────────────────
    heading1(doc, "6. Non-Functional Requirements (NFR) — Final State")
    make_table(doc,
        headers=["NFR Category", "Requirement"],
        rows=[
            ["Security",
             "Passwords are never stored or compared in plain text — SHA-256 hashing is applied at all registration and login points via PasswordUtil.hash(). Session state is held in a singleton that is explicitly cleared on logout. No credentials persist beyond the session. RBAC guards are applied at the controller level before every instructor-only action."],
            ["Performance",
             "Gradebook queries (class-wide grid) execute in a single JOIN (fetchGradebookForCourse) rather than N+1 queries per student. All primary views target loading within 2–3 seconds on typical lab hardware at class sizes up to ~200 students per course."],
            ["Reliability",
             "Grade saves are atomic — Submission.saveGrade() issues a single parameterised UPDATE; no partial grade state is possible. Session logout completes synchronously before the LoginView is displayed. Submission timestamps are set server-side, not client-side, ensuring consistent status evaluation."],
            ["Usability",
             "Role-adaptive UI: instructor-only buttons (Post Announcement, Create Assignment, Grade Submissions) are conditionally rendered — not merely disabled — for students. Status badges (Submitted / Late / Missing / Not Submitted) use both colour and text so they are never colour-only. Gradebook cells show — for ungraded rather than blank or null."],
            ["Maintainability",
             "All Sprint 3 features follow the established layered MVC architecture (controller / model / view). RBAC checks, grade validation, and gradebook queries live in model/controller classes, not in FXML event handlers. All new views follow the same FXML + CSS structure as existing ones."],
            ["Auditability",
             "Graded submissions retain the original submission timestamp alongside the new grade and feedback fields — no information is overwritten. Session login is gated on credential verification; no anonymous access path remains in the application after Sprint 3."],
        ]
    )

    # ── 7. Testing ────────────────────────────────────────────────────────────
    heading1(doc, "7. Testing and Quality Assurance")

    heading2(doc, "7.1 Test Plan Summary")
    add_body_para(doc,
        "Testing for Sprint 3 targeted four risk areas: authentication correctness (credentials must be verified, "
        "never bypassed), RBAC integrity (students must never reach instructor actions), grade validation (numeric "
        "bounds must be enforced before DB write), and submission status accuracy (all four status paths must be "
        "reachable and correct). The seeded database (`make seed`) provided a deterministic, repeatable test state "
        "for all manual test runs. See Part B Section 6 for the full test plan, test cases, white-box coverage "
        "analysis, and bug report."
    )

    heading2(doc, "7.2 Black-Box Test Cases (Summary)")
    add_body_para(doc,
        "25 black-box test cases were designed across four modules using equivalence partitioning, boundary value "
        "analysis, and error guessing. All 25 passed. Key coverage:"
    )
    add_bullet(doc, "Login/Register: valid credentials, wrong password, unknown email, empty fields, SQL injection attempt, duplicate username/email")
    add_bullet(doc, "RBAC: student blocked from instructor actions, enrollment filtering correct for both roles")
    add_bullet(doc, "Grading: valid grades, grade = 0, grade = total_points, grade > total_points, non-numeric input, empty feedback")
    add_bullet(doc, "Submission status: all four states (Submitted, Late, Missing, Not Submitted) verified against seeded data")

    heading2(doc, "7.3 White-Box Testing")
    add_body_para(doc,
        "Branch coverage analysis was performed on `LoginController.handleLogin()` (4/4 branches, 100%) and "
        "`GradingController.handleSave()` (4/4 branches, 100%). Statement coverage of all reachable paths in "
        "`Assignment.fetchWithStatusForUser()` was achieved via TC-21 through TC-24. "
        "See Part B Section 6c for full coverage tables."
    )

    heading2(doc, "7.4 Bug Report Summary")
    make_table(doc,
        headers=["Bug ID", "Module", "Severity", "Status"],
        rows=[
            ["BUG-01", "Grading",     "P1 — NullPointerException on empty submission list", "Fixed"],
            ["BUG-02", "RBAC",        "P1 — Instructor saw all courses, not just their own", "Fixed"],
            ["BUG-03", "Gradebook",   "P2 — GridPane column misalignment with missing submissions", "Fixed"],
            ["BUG-04", "Attachments", "P1 — Attachment chips not rendering after stream refresh", "Fixed"],
        ]
    )
    add_body_para(doc,
        "All four bugs were identified during manual test execution and fixed before submission. "
        "No unfixed P1 bugs remain."
    )

    heading2(doc, "7.5 Test Results and Observations")
    add_body_para(doc,
        "All 25 black-box test cases passed. All three seeded accounts (ali, ibbi, sara) were used to verify "
        "role-specific behaviour end-to-end. The SQL injection test (TC-06) confirmed that parameterised queries "
        "prevent injection at the login boundary. Boundary tests on the grading validator (TC-16 to TC-18) confirmed "
        "that grade = 0 and grade = total_points are both accepted while grade = total_points + 1 is correctly rejected."
    )

    heading2(doc, "7.6 How Testing Improved System Quality")
    add_body_para(doc,
        "Testing directly uncovered four bugs that would have been user-facing in submission. BUG-01 would have "
        "crashed the Grading modal silently for any assignment with no submissions yet — a common early-sprint state. "
        "BUG-02 broke the core RBAC invariant entirely, allowing instructors to see courses they did not own. BUG-04 "
        "made the entire attachment feature invisible to students after posting. All were caught in the test phase and "
        "fixed before the final release. The structured test plan also forced explicit coverage of all four submission "
        "status paths, which revealed a gap in the status derivation logic for the Not Submitted case that was "
        "corrected before integration."
    )

    # ── 8. Screenshots ────────────────────────────────────────────────────────
    heading1(doc, "8. Sprint 3 Implementation Screenshots")

    screenshots = [
        ("[Screenshot 1 — Login view with Login and Register tabs]",
         "LoginView — tabbed interface with Login (email + password + Sign In) and Register "
         "(username + email + password + role picker + Register) tabs. Inline error labels appear on validation failures."),
        ("[Screenshot 2 — Course Catalog filtered by role]",
         "CourseCatalog — students see only courses they are enrolled in; instructors see courses they created. "
         "Log Out button in the header clears the session."),
        ("[Screenshot 3 — Announcement with file and link attachment chips]",
         "Course Stream — announcement card with a file attachment chip and a URL attachment chip rendered below the body text."),
        ("[Screenshot 4 — Assignments tab: student view with status badges]",
         "Assignments Tab (student) — assignment cards showing deadline, point value, late-policy badge, and "
         "colour-coded status badge (Submitted / Late / Missing / Not Submitted)."),
        ("[Screenshot 5 — Assignments tab: instructor view with Grade Submissions button]",
         'Assignments Tab (instructor) — same card layout; "Grade Submissions" button replaces the status badge '
         "and opens the Grading modal."),
        ("[Screenshot 6 — Grading modal]",
         "GradingView — two-panel modal: left ListView of submissions (student name + status), right panel with "
         "numeric grade field, feedback TextArea, and Save button. Pre-filled for already-graded submissions."),
        ("[Screenshot 7 — Student gradebook]",
         "Student Gradebook — personal grade table showing assignment name, score out of total points, and "
         "instructor feedback per row."),
        ("[Screenshot 8 — Instructor class gradebook]",
         'Instructor Class Gradebook — GridPane with student names as rows and assignment titles as columns; '
         'numeric grades or "—" for ungraded submissions.'),
        ("[Screenshot 9 — People tab]",
         "People Tab — course roster cards with avatar initials, username, email, and role badge; "
         "instructors grouped above students."),
    ]

    for placeholder, caption in screenshots:
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.space_before = Pt(6)
        p_ph.paragraph_format.space_after  = Pt(1)
        rph = p_ph.add_run(placeholder)
        rph.font.name   = "Times New Roman"
        rph.font.size   = Pt(10)
        rph.font.italic = True
        rph.font.color.rgb = GRAY

        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after  = Pt(8)
        p_cap.paragraph_format.left_indent  = Inches(0.2)
        rc = p_cap.add_run("Caption: " + caption)
        rc.font.name  = "Times New Roman"
        rc.font.size  = Pt(10)

    # ── 8b. Work Division ─────────────────────────────────────────────────────
    heading1(doc, "8. Work Division for Sprint 3")
    make_table(doc,
        headers=["Member", "Sprint 3 Responsibilities"],
        rows=[
            ["Ibraheem Farooq\n(PM / Scrum Master / Developer)",
             "Sprint 3 planning and Trello board management across all sprint stages. Branch and PR oversight — "
             "reviewed and merged feature/auth, feature/enrollment-rbac, and feature/announcements-materials PRs. "
             "Implemented end-to-end authentication and RBAC: tabbed Login/Register UI, Session singleton, "
             "PasswordUtil, LoginController, enrollment-based catalog filtering, and Access Denied guard on "
             "instructor-only actions. Implemented announcement file and link attachments (US 12 & 13): "
             "CreateAnnouncement extension, announcement_attachments DB table, and attachment chip rendering in "
             "the course stream. Final review and packaging of deliverable artefacts."],
            ["Wajih-Ur-Raza Asif\n(Requirements Architect / Tester / Developer)",
             "Sprint 3 user stories, sub-story breakdowns, and all six structured specifications. Consistency "
             "verification between acceptance criteria and actual implemented behaviour. Deliverable 3 document "
             "authoring and completeness review. Implemented the Assignments tab and People tab (US 15 & 16): "
             "Assignment.fetchWithStatusForUser(), status badge rendering, Grade Submissions button, "
             "Course.fetchEnrolledUsers(), and roster card rendering. Implemented the full grading pipeline "
             "(US 19–22): grade/feedback DB schema, Submission model extensions (fetchByAssignmentId, saveGrade, "
             "fetchGradedForStudent, fetchGradebookForCourse), GradingView.fxml + GradingController, and "
             "GradebookView.fxml + GradebookController (student personal view and instructor class-wide grid). "
             "End-to-end test plan execution using seeded test data."],
            ["Ali Aan Khowaja\n(Developer / UI Designer)",
             "Lead implementation of all Sprint 2 features: Announcements, Assignment Creation, AssignmentDetails "
             "with file submission flow and threaded comments, and the To-Do List. UI design and screen layout for "
             "all Sprint 2 and Sprint 3 views. Produced Sprint 3 implementation screenshots."],
        ]
    )

    # ── 9. Retrospective ─────────────────────────────────────────────────────
    heading1(doc, "9. Retrospective")

    heading2(doc, "What went well")
    wents = [
        "All Sprint 3 user stories were completed before the submission deadline, including the two additions "
        "(Assignments tab, People tab) that were not in the original Sprint 2 backlog carry-forward.",
        "The layered MVC architecture established in Sprint 1 absorbed all new features without requiring "
        "structural changes to existing code. The grading pipeline (DB schema extension to model methods to "
        "controller to FXML view) followed the same pattern as every prior feature.",
        "The idempotent seed.sql with pre-configured accounts (ali, ibbi, sara) and realistic submission states "
        "— including a mix of graded and ungraded submissions across two courses — significantly accelerated "
        "manual testing and made it easy to verify both student and instructor views without manually creating data.",
        "The branch-per-feature strategy with descriptive PR titles kept code review manageable. The one merge "
        "conflict encountered (feature/assignments-tab diverging from main after grading branches merged) was "
        "resolved cleanly.",
    ]
    for w in wents:
        add_bullet(doc, w)

    heading2(doc, "What could be improved")
    improvements = [
        "Authentication was implemented in Sprint 3 when it should have been Sprint 1. All D1 and D2 features "
        "operated without any session guard, which meant RBAC had to be retroactively applied to every existing "
        "controller in Sprint 3. This added scope to an already dense sprint.",
        "The grading feature chain (grading-db to grading-ui to gradebook) created a linear merge dependency. "
        "The gradebook branch could not be started until grading-ui merged, which serialised work that could "
        "otherwise have run in parallel.",
        "US 15–16 (Assignments tab, People tab) were not in any prior backlog despite being obvious complements "
        "to the D2 submission workflow. Identifying them earlier would have allowed them to be built during D2 "
        "when the assignment infrastructure was fresh.",
    ]
    for imp in improvements:
        add_bullet(doc, imp)

    heading2(doc, "What we would do differently")
    diffs = [
        "Implement user registration and login in Sprint 1 so that all subsequent features are built on top of "
        "an authenticated, role-aware session from the start.",
        "Identify course-home navigation features (Assignments tab, People tab) during D2 planning and include "
        "them in the D2 sprint alongside the assignment and submission work they depend on.",
        "Use an integration branch (feature to integration to main) to allow the grading feature chain to be "
        "worked on in parallel with less merge risk.",
    ]
    for d in diffs:
        add_bullet(doc, d)

    # ── 11. Final System Summary ──────────────────────────────────────────────
    heading1(doc, "11. Final System Summary")
    add_body_para(doc,
        "Naawbi LMS was built across three sprints, each delivering a distinct layer of a complete "
        "learning management system."
    )
    add_body_para(doc,
        "**Sprint 1** established the structural foundation: instructors can create courses with a unique class "
        "code, and students can enrol using that code. The course catalogue gives each user a personalised home "
        "for their courses."
    )
    add_body_para(doc,
        "**Sprint 2** added the content and submission layer: instructors can post announcements and create "
        "assignments with configurable deadlines and late policies; students can upload file submissions which "
        "are timestamped and automatically tagged as Submitted, Late, or Missing. A student-facing To-Do list "
        "aggregates pending work across all enrolled courses with urgency indicators. Threaded comments on "
        "submissions allow instructor-student communication within a submission context."
    )
    add_body_para(doc,
        "**Sprint 3** closed the platform: every user now registers with a verified identity and role-specific "
        "session. Role-based access control ensures students see only their enrolled courses and cannot reach "
        "instructor actions. Instructors can enrich announcements with file and link attachments. The Assignments "
        "tab gives users a per-course view of all assignments and their status without leaving the course home; "
        "the People tab shows the full enrolled roster. A grading interface lets instructors assign numeric grades "
        "and feedback to individual submissions, and both a student personal gradebook and an instructor "
        "class-wide gradebook surface those results."
    )
    make_table(doc,
        headers=["Sprint", "Deliverable", "Core Contribution"],
        rows=[
            ["Sprint 1", "D1", "Course structure — creation, catalogue, class-code enrolment"],
            ["Sprint 2", "D2", "Content and submissions — announcements, assignments, file uploads, status tracking, To-Do list"],
            ["Sprint 3", "D3", "Identity and evaluation — authentication, RBAC, attachments, assignments/people tabs, grading, gradebook"],
        ]
    )
    add_body_para(doc,
        "The final system fulfils its original vision: a clean, course-centric LMS where instructors manage "
        "content and evaluations, and students track progress and performance, without requiring any external "
        "tool for any core academic workflow. All planned user stories across all three sprints have been "
        "delivered and verified."
    )


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Default style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Section 0 — cover (no header/footer on cover)
    section0 = doc.sections[0]
    set_doc_margins(section0, top=1.2, bottom=1.0, left=1.2, right=1.0)
    section0.different_first_page_header_footer = True

    build_cover(doc)
    build_toc(doc)

    # Start a new section for body pages so we can add header/footer
    # (the cover page break we added already moves us, but we need a
    #  section break to switch header on)
    # Add section break (continuous to next page) before body
    # Actually the page break after cover is enough; we just enable
    # header starting from page 2 via different_first_page.

    add_header_footer(section0, doc)

    build_part_a(doc)
    build_part_b(doc)
    build_part_c(doc)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
